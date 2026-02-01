"""File format handlers for reading various source material types."""

import csv
import io
import json
import logging
from dataclasses import dataclass, field
from pathlib import Path

logger = logging.getLogger(__name__)


@dataclass
class RawDocument:
    """A raw document read from a source file."""

    content: str
    source_path: str
    file_type: str
    metadata: dict = field(default_factory=dict)


def read_text_file(path: Path) -> RawDocument:
    """Read a plain .txt or .md file."""
    content = path.read_text(encoding="utf-8")
    return RawDocument(
        content=content,
        source_path=str(path),
        file_type=path.suffix.lstrip("."),
    )


def read_csv_file(path: Path) -> RawDocument:
    """Read a .csv vocabulary list.

    Expects columns like: hindi, english, romanization, notes
    Returns the CSV data as a JSON string for downstream processing.
    """
    content = path.read_text(encoding="utf-8")
    reader = csv.DictReader(io.StringIO(content))
    rows = list(reader)
    return RawDocument(
        content=json.dumps(rows, ensure_ascii=False),
        source_path=str(path),
        file_type="csv",
        metadata={"row_count": len(rows), "columns": list(rows[0].keys()) if rows else []},
    )


def read_json_file(path: Path) -> RawDocument:
    """Read a .json structured export."""
    content = path.read_text(encoding="utf-8")
    # Validate it's valid JSON
    data = json.loads(content)
    return RawDocument(
        content=json.dumps(data, ensure_ascii=False),
        source_path=str(path),
        file_type="json",
        metadata={"item_count": len(data) if isinstance(data, list) else 1},
    )


def read_docx_file(path: Path) -> RawDocument:
    """Read a .docx file (Google Docs export).

    Requires the python-docx package.
    """
    try:
        import docx
    except ImportError:
        logger.warning("python-docx not installed, skipping %s", path)
        raise ImportError("Install python-docx to process .docx files: pip install python-docx")

    doc = docx.Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    content = "\n".join(paragraphs)
    return RawDocument(
        content=content,
        source_path=str(path),
        file_type="docx",
        metadata={"paragraph_count": len(paragraphs)},
    )


# Map file extensions to their handler functions
HANDLERS: dict[str, callable] = {
    ".txt": read_text_file,
    ".md": read_text_file,
    ".csv": read_csv_file,
    ".json": read_json_file,
    ".docx": read_docx_file,
}

SUPPORTED_EXTENSIONS = set(HANDLERS.keys())


def read_file(path: Path) -> RawDocument:
    """Read a file using the appropriate handler based on extension."""
    ext = path.suffix.lower()
    handler = HANDLERS.get(ext)
    if handler is None:
        raise ValueError(f"Unsupported file type: {ext}. Supported: {SUPPORTED_EXTENSIONS}")
    logger.info("Reading %s (%s)", path.name, ext)
    return handler(path)


def read_directory(directory: Path) -> list[RawDocument]:
    """Recursively read all supported files from a directory."""
    documents = []
    for ext in SUPPORTED_EXTENSIONS:
        for path in sorted(directory.rglob(f"*{ext}")):
            try:
                documents.append(read_file(path))
            except Exception:
                logger.exception("Failed to read %s", path)
    logger.info("Read %d documents from %s", len(documents), directory)
    return documents
